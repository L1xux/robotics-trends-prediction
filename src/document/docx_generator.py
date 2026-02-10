"""
DOCX Generator

Markdown 보고서를 DOCX 파일로 변환
"""

from typing import Optional
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


class DocxGenerator:
    """
    Markdown to DOCX Converter
    
    Simple and robust conversion with minimal formatting
    """
    
    def __init__(self):
        """Initialize DOCX Generator"""
        pass
    
    def generate(
        self,
        markdown_content: str,
        output_path: str,
        title: Optional[str] = None
    ) -> str:
        """
        Generate DOCX from markdown
        
        Args:
            markdown_content: Markdown text
            output_path: Output file path
            title: Document title (optional)
        
        Returns:
            Path to generated DOCX file
        """
        try:
            # Create document
            doc = Document()
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Malgun Gothic'  # Korean support
            font.size = Pt(11)
            
            # Process markdown line by line
            lines = markdown_content.split('\n')
            
            for line in lines:
                line = line.rstrip()
                
                # Skip empty lines
                if not line:
                    continue
                
                # H1 (# Title)
                if line.startswith('# '):
                    text = line[2:].strip()
                    p = doc.add_heading(text, level=1)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # H2 (## Section)
                elif line.startswith('## '):
                    text = line[3:].strip()
                    doc.add_heading(text, level=2)
                
                # H3 (### Subsection)
                elif line.startswith('### '):
                    text = line[4:].strip()
                    doc.add_heading(text, level=3)
                
                # Horizontal rule (---)
                elif line.startswith('---'):
                    doc.add_paragraph('_' * 50)
                
                # Bold/Italic/Normal text
                else:
                    # Remove markdown bold/italic for simplicity
                    text = self._clean_markdown(line)
                    if text:
                        doc.add_paragraph(text)
            
            # Save
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            
            print(f"DOCX generated: {output_path}")
            
            return str(output_path)
        
        except Exception as e:
            print(f"DOCX generation failed: {e}")
            raise
    
    def _clean_markdown(self, text: str) -> str:
        """
        Remove markdown formatting for simple text
        
        Args:
            text: Markdown text
        
        Returns:
            Plain text
        """
        # Remove bold
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        # Remove italic
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        # Remove inline code
        text = re.sub(r'`(.+?)`', r'\1', text)
        
        return text.strip()


def generate_docx(
    markdown_content: str,
    output_path: str,
    title: Optional[str] = None
) -> str:
    """
    Convenience function to generate DOCX
    
    Args:
        markdown_content: Markdown text
        output_path: Output file path
        title: Document title (optional)
    
    Returns:
        Path to generated DOCX file
    """
    generator = DocxGenerator()
    return generator.generate(markdown_content, output_path, title)
